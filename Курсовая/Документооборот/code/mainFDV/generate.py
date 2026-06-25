# Создайте тестовый шаблон программно
from docx import Document
from docxtpl import DocxTemplate
import os

# Создаем документ
doc = Document()
doc.add_paragraph("{{ familia }} {{ name }} {{ otchestvo }}")
doc.add_paragraph("Тип практики: {{ tip }}")
doc.add_paragraph("Группа: {{ group }}")

# Путь для сохранения на Linux/Mac
# Вариант 1: Рабочий стол
save_path = '/tmp/test_template.docx'

# Вариант 2: /tmp (временная папка)
# save_path = '/tmp/test_template.docx'

# Вариант 3: Текущая папка
# save_path = 'test_template.docx'

# Сохраняем файл
doc.save(save_path)
print(f"Файл сохранен: {save_path}")